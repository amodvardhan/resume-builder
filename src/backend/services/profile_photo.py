"""User profile photo: validation, storage, and path resolution for resume exports."""

from __future__ import annotations

import io
import logging
import uuid
from pathlib import Path

from PIL import Image

from src.backend.config import settings

logger = logging.getLogger(__name__)

_MAX_BYTES = 4 * 1024 * 1024
_MAX_DIMENSION = 1024


def _detect_format(header: bytes) -> str | None:
    if len(header) >= 3 and header[:3] == b"\xff\xd8\xff":
        return "jpeg"
    if len(header) >= 8 and header[:8] == b"\x89PNG\r\n\x1a\n":
        return "png"
    if len(header) >= 12 and header[:4] == b"RIFF" and header[8:12] == b"WEBP":
        return "webp"
    return None


def validate_process_and_save_profile_photo(
    raw: bytes,
    *,
    user_id: uuid.UUID,
) -> str:
    """Validate image bytes, normalize with Pillow, save under ``profile_photos_dir``.

    Returns the absolute filesystem path stored on the user row.
    """
    if len(raw) > _MAX_BYTES:
        raise ValueError("Image exceeds maximum size (4 MB)")

    fmt = _detect_format(raw[:32])
    if fmt is None:
        raise ValueError("Only JPEG, PNG, or WebP images are allowed")

    settings.profile_photos_dir.mkdir(parents=True, exist_ok=True)
    dest = settings.profile_photos_dir / f"{user_id}.jpg"

    try:
        im = Image.open(io.BytesIO(raw))
        im = im.convert("RGB")
        im.thumbnail((_MAX_DIMENSION, _MAX_DIMENSION), Image.Resampling.LANCZOS)
        im.save(dest, format="JPEG", quality=90, optimize=True)
    except Exception as exc:
        logger.warning("Profile photo processing failed: %s", exc)
        raise ValueError("Could not process image — try another file") from exc

    return str(dest.resolve())


def resolved_photo_path(user_profile_photo_path: str | None) -> Path | None:
    """Return a readable path for tailoring, or None if unset / missing."""
    if not user_profile_photo_path or not str(user_profile_photo_path).strip():
        return None
    p = Path(user_profile_photo_path).expanduser()
    if not p.is_file():
        return None
    return p


def prepend_photo_to_docx_file(docx_path: Path, photo_path: Path) -> None:
    """Insert a right-aligned headshot at the top of an existing .docx (user template path)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    if not photo_path.is_file():
        return

    doc = Document(str(docx_path))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.add_picture(str(photo_path), width=Inches(1.15))
    p.paragraph_format.space_after = Pt(10)
    body = doc.element.body
    body.insert(0, p._element)
    doc.save(str(docx_path))
