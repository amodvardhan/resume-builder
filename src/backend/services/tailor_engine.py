"""
Tailoring engine — LangChain orchestration + python-docx injection.

Responsibilities:
  1. Accept a JD, user profile, and template reference.
  2. Build a LangChain chain (ChatPromptTemplate → ChatOpenAI → PydanticOutputParser).
  3. Inject the parsed JSON fields into a *copy* of the .docx template.
  4. Return the output file path and generated cover-letter text.
"""

from __future__ import annotations

import logging
import shutil
import uuid
from pathlib import Path
from typing import Any

from docx import Document
from langchain_core.output_parsers import PydanticOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
from pydantic import BaseModel, Field, field_validator

from src.backend.config import settings

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _coerce_experience_to_str(value: Any) -> str:
    """Accept a plain string or a dict the LLM may produce and flatten it."""
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        parts: list[str] = []
        if value.get("title"):
            line = str(value["title"])
            if value.get("organization"):
                line += f" | {value['organization']}"
            if value.get("dates"):
                line += f" | {value['dates']}"
            parts.append(line)
        for bullet in value.get("achievements", value.get("bullets", [])):
            parts.append(f"• {bullet}")
        if value.get("description"):
            parts.append(str(value["description"]))
        return "\n".join(parts) if parts else str(value)
    return str(value)


# ---------------------------------------------------------------------------
# Pydantic model whose fields map 1:1 to {{TAGS}} inside .docx templates
# ---------------------------------------------------------------------------


class TailoredContent(BaseModel):
    """Structured output the LLM *must* produce.

    Uses a dynamic list of experiences so ALL roles from the candidate's
    resume are captured — not just 3.
    """

    summary: str = Field(
        description=(
            "A 4-5 sentence professional summary that bridges the candidate's "
            "REAL background (extracted from their resume) with the target JD. "
            "Must reference actual roles, achievements, and skills from the "
            "resume. Be detailed and specific — this is the executive overview."
        )
    )
    experiences: list[str] = Field(
        description=(
            "A list of EVERY SINGLE job/role entry found in the candidate's "
            "resume. You MUST include ALL roles — if the resume has 5 roles "
            "you must return 5 entries, if it has 8 you must return 8. "
            "NEVER drop or skip any role. Order by relevance to the target JD. "
            "Each entry MUST be a single flat string formatted as: "
            "'Title | Organization | Dates\\n• achievement 1\\n• achievement 2\\n…'. "
            "Include 4-7 DETAILED achievement bullets per role — each bullet "
            "should be a substantive sentence with impact, metrics, or outcomes. "
            "MUST use REAL titles, organizations, dates, and achievements from "
            "the resume. NEVER fabricate roles or achievements. "
            "If a role has many bullets in the resume, keep them ALL."
        )
    )
    skills: str = Field(
        description=(
            "A COMPREHENSIVE comma-separated list of ALL skills from the "
            "candidate's resume, reordered by relevance to the JD. Include "
            "technical skills, tools, frameworks, methodologies, languages, "
            "and soft skills. Do NOT truncate — include every skill mentioned."
        )
    )
    education: str = Field(
        description=(
            "ALL education entries from the candidate's resume. Include "
            "every degree, institution, date, honours, GPA (if present), and "
            "relevant coursework. If there are multiple degrees, include ALL. "
            "If no education section exists, output 'Not specified in resume'."
        )
    )
    certifications: str = Field(
        default="",
        description=(
            "ALL certifications, licenses, professional development entries, "
            "and training programs from the candidate's resume. If none found, "
            "return an empty string."
        )
    )

    @field_validator("experiences", mode="before")
    @classmethod
    def _flatten_experiences(cls, v: Any) -> list[str]:
        if isinstance(v, list):
            return [_coerce_experience_to_str(item) for item in v]
        return [_coerce_experience_to_str(v)]


class CoverLetterContent(BaseModel):
    """Structured cover-letter output."""

    cover_letter: str = Field(description="Full cover-letter body text")


# ---------------------------------------------------------------------------
# Prompt templates
# ---------------------------------------------------------------------------

_SENTIMENT_PREAMBLES: dict[str, str] = {
    "formal": (
        "You write in a polished, formal register suitable for government and "
        "intergovernmental organizations. Avoid colloquialisms. Maintain "
        "gravitas while letting genuine professional passion come through. "
        "The letter should feel authoritative yet approachable — written by "
        "a human being, not a template generator."
    ),
    "mission-driven": (
        "You write with genuine passion about the organization's mission. "
        "Show deep alignment between the candidate's values and the org's "
        "purpose. Use emotive, compelling language that connects the "
        "candidate's personal motivations to the organization's impact. "
        "The reader should feel the candidate truly cares about the cause, "
        "not just the job title. Weave in specific examples from their "
        "background that demonstrate this alignment authentically."
    ),
    "conversational": (
        "You write in a warm yet professional tone that feels personable "
        "while remaining appropriate for a corporate setting. The letter "
        "should read as though a real person wrote it with care — not "
        "a machine. Use natural transitions, occasionally address the "
        "reader, and let the candidate's personality shine through. Avoid "
        "stiff constructions; prefer active voice and vivid language."
    ),
}

_RESUME_SYSTEM = (
    "You are a senior resume-tailoring specialist. Your job has TWO phases:\n"
    "  Phase A — EXTRACT verbatim facts from the candidate's attached resume.\n"
    "  Phase B — TAILOR the extracted content to align with the target job description.\n\n"
    "The candidate's core skills are: {core_skills}.\n"
    "{history_context}"
    "The candidate's ACTUAL resume (uploaded by them) is below. "
    "This is the ONLY source of truth for their background:\n"
    "--- BEGIN CANDIDATE RESUME ---\n"
    "{resume_text}\n"
    "--- END CANDIDATE RESUME ---\n\n"
    "═══════════════════════════════════════════════════\n"
    "PHASE A — EXTRACTION RULES (MANDATORY):\n"
    "═══════════════════════════════════════════════════\n"
    "1. **EXPERIENCE**: You MUST find and extract EVERY SINGLE job/role "
    "entry from the resume. Read the ENTIRE resume from start to end. "
    "Look for sections marked [SECTION: EXPERIENCE], [ROLE: ...], or any "
    "text describing a job position (title + company + dates). "
    "Count the total number of roles. If the resume has 3 roles, extract 3. "
    "If it has 7 roles, extract 7. NEVER skip a role. "
    "For each role extract: REAL job title, REAL organization name, REAL "
    "dates, and ALL achievement bullets EXACTLY as they appear. "
    "Do NOT invent, fabricate, or hallucinate any role, company, date, or "
    "achievement not in the resume.\n"
    "2. **SKILLS**: Find the skills/competencies section (look for "
    "[SECTION: SKILLS] or similar). Extract ALL skills listed. If no "
    "explicit skills section, infer skills from the experience descriptions. "
    "Be comprehensive — include every technical skill, tool, framework, "
    "methodology, and soft skill.\n"
    "3. **EDUCATION**: Find ALL education entries (look for "
    "[SECTION: EDUCATION] or similar). Extract EVERY degree, institution, "
    "date, honours, GPA. If multiple degrees exist, include ALL of them. "
    "If no education section exists, output 'Not specified in resume'.\n"
    "4. **CERTIFICATIONS**: Find ALL certifications, licenses, training "
    "programs (look for [SECTION: CERTIFICATIONS] or similar). Extract "
    "verbatim. If none exist, return empty string.\n\n"
    "═══════════════════════════════════════════════════\n"
    "PHASE B — TAILORING RULES:\n"
    "═══════════════════════════════════════════════════\n"
    "5. **SUMMARY**: Write a 4-5 sentence professional summary bridging "
    "the candidate's ACTUAL background with the target JD. Reference "
    "specific roles and achievements. Do NOT claim expertise the resume "
    "does not support.\n"
    "6. **EXPERIENCES**: For EVERY extracted role (ALL of them — do NOT "
    "drop any), keep the REAL title, org, and dates. Rephrase the "
    "achievement bullets to emphasize JD alignment. You may reword for "
    "clarity but NEVER add fabricated achievements. Include 4-7 detailed "
    "bullets per role. Order roles by relevance to the JD.\n"
    "7. **SKILLS**: Reorder extracted skills by JD relevance. You may add "
    "skills clearly demonstrated in experience but NEVER add unsupported "
    "skills.\n"
    "8. **EDUCATION**: Return ALL education entries exactly as extracted. "
    "Do not alter degrees, institutions, or dates.\n\n"
    "═══════════════════════════════════════════════════\n"
    "CRITICAL LENGTH REQUIREMENTS:\n"
    "═══════════════════════════════════════════════════\n"
    "- The output MUST be COMPREHENSIVE. Do NOT summarise or abbreviate.\n"
    "- The `experiences` array MUST have one entry per role found in the "
    "resume. If the resume has 5 roles, the array MUST have 5 elements.\n"
    "- Each experience entry MUST have 4-7 substantive achievement bullets "
    "with impact, metrics, and outcomes where available.\n"
    "- The resume may span TWO PAGES. Do NOT shorten to fit one page.\n"
    "- `skills` must be an exhaustive comma-separated list.\n"
    "- `education` must include ALL degrees and qualifications.\n\n"
    "═══════════════════════════════════════════════════\n"
    "OUTPUT FORMAT RULES:\n"
    "═══════════════════════════════════════════════════\n"
    "9. **`experiences` is a JSON array of strings.** Each element is one "
    "role: \"Title | Organization | Dates\\n• bullet 1\\n• bullet 2…\". "
    "Use \\n for line breaks within each entry.\n"
    "10. `skills`: single comma-separated string, comprehensive.\n"
    "11. `education`: single string with ALL degrees, institutions, dates.\n"
    "12. `certifications`: single string, empty if none.\n"
    "13. `summary`: 4-5 detailed sentences.\n\n"
    "{format_instructions}"
)

_RESUME_HUMAN = (
    "Job title: {job_title}\n"
    "Organization: {organization}\n\n"
    "Job description:\n{job_description}"
)

_COVER_LETTER_SYSTEM = (
    "{sentiment_preamble}\n\n"
    "You are a senior career strategist writing a compelling, detailed cover "
    "letter for a candidate.\n"
    "Candidate core skills: {core_skills}.\n"
    "They are applying for {job_title} at {organization}.\n\n"
    "The candidate's ACTUAL resume (uploaded by them) is below:\n"
    "--- BEGIN CANDIDATE RESUME ---\n"
    "{resume_text}\n"
    "--- END CANDIDATE RESUME ---\n\n"
    "═══════════════════════════════════════════════════\n"
    "COVER LETTER RULES:\n"
    "═══════════════════════════════════════════════════\n"
    "1. **LENGTH**: Write 5-7 substantial paragraphs. This is a DETAILED "
    "cover letter, not a brief note. Each paragraph should serve a "
    "distinct purpose.\n"
    "2. **STRUCTURE**:\n"
    "   - Opening: A compelling hook that connects the candidate to the "
    "     role and organization. Show genuine enthusiasm.\n"
    "   - Body (2-3 paragraphs): Deep-dive into SPECIFIC, REAL achievements "
    "     from the resume. For each achievement, explain the CONTEXT, the "
    "     candidate's ACTION, and the RESULT/IMPACT. Tie each example "
    "     directly to a requirement from the JD.\n"
    "   - Values/Mission paragraph: Connect the candidate's professional "
    "     values and motivations to the organization's mission and culture.\n"
    "   - Forward-looking paragraph: Articulate what the candidate would "
    "     bring to the role and team, grounded in their real experience.\n"
    "   - Closing: Confident, warm close with a call to action.\n"
    "3. **HUMANISATION**: The letter MUST read as though a thoughtful "
    "human being wrote it. Use natural transitions, varied sentence "
    "structure, and occasional personal reflection. Avoid robotic "
    "phrases like 'I am writing to express my interest' or 'I believe "
    "I would be a great fit'. Instead, open with something specific "
    "and genuine.\n"
    "4. **SPECIFICITY**: Reference REAL job titles, organizations, metrics, "
    "project names, and accomplishments from the resume. Vague claims "
    "are forbidden.\n"
    "5. **NEVER** claim the candidate has experience or achievements not "
    "in their resume.\n"
    "6. Produce a COMPLETE, ready-to-send cover letter body (no "
    "placeholders, no [Name] tokens).\n\n"
    "{format_instructions}"
)

_COVER_LETTER_HUMAN = "Job description:\n{job_description}"


# ---------------------------------------------------------------------------
# Document injection
# ---------------------------------------------------------------------------


def _replace_tags_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> None:
    """Replace {{TAG}} placeholders that may span multiple XML runs."""
    full_text = "".join(run.text for run in paragraph.runs)
    if not any(("{{" + tag.upper() + "}}") in full_text for tag in replacements):
        return

    for tag, value in replacements.items():
        placeholder = "{{" + tag.upper() + "}}"
        full_text = full_text.replace(placeholder, value)

    if not paragraph.runs:
        return

    paragraph.runs[0].text = full_text
    for run in paragraph.runs[1:]:
        run.text = ""


def inject_into_template(template_path: Path, replacements: dict[str, str]) -> str:
    """
    Copy the master template and replace every ``{{TAG}}`` with the matching
    value from *replacements*.  Returns the output **filename** (not full path).
    """
    settings.output_dir.mkdir(parents=True, exist_ok=True)
    output_filename = f"{uuid.uuid4()}.docx"
    output_path = settings.output_dir / output_filename

    shutil.copy2(template_path, output_path)
    doc = Document(str(output_path))

    for paragraph in doc.paragraphs:
        _replace_tags_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_tags_in_paragraph(paragraph, replacements)

    doc.save(str(output_path))
    return output_filename


# ---------------------------------------------------------------------------
# From-scratch document builders (when no user template is uploaded)
# ---------------------------------------------------------------------------

def _add_heading(doc: Any, text: str, level: int = 2) -> None:
    from docx.shared import Pt, RGBColor
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        run.font.size = Pt(12 if level == 2 else 11)


def build_resume_docx(content: dict[str, Any]) -> str:
    """Build a polished resume docx from scratch — no template required."""
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    settings.output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    if content.get("summary"):
        _add_heading(doc, "Professional Summary")
        doc.add_paragraph(content["summary"])

    experiences = content.get("experiences", [])
    if experiences:
        _add_heading(doc, "Professional Experience")
        for exp_text in experiences:
            lines = exp_text.split("\n")
            for j, line in enumerate(lines):
                p = doc.add_paragraph()
                stripped = line.strip()
                if j == 0:
                    run = p.add_run(stripped)
                    run.bold = True
                    run.font.size = Pt(10.5)
                elif stripped.startswith("•"):
                    p.style = doc.styles.get("List Bullet", doc.styles["Normal"])
                    p.add_run(stripped.lstrip("• "))
                else:
                    p.add_run(stripped)

    if content.get("skills"):
        _add_heading(doc, "Skills")
        doc.add_paragraph(content["skills"])

    if content.get("education"):
        _add_heading(doc, "Education")
        doc.add_paragraph(content["education"])

    if content.get("certifications"):
        _add_heading(doc, "Certifications")
        doc.add_paragraph(content["certifications"])

    output_filename = f"{uuid.uuid4()}.docx"
    output_path = settings.output_dir / output_filename
    doc.save(str(output_path))
    return output_filename


def build_cover_letter_docx(cover_letter_text: str) -> str:
    """Build a standalone cover-letter docx."""
    from docx.shared import Pt, Inches, RGBColor

    settings.output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.15)
        section.right_margin = Inches(1.15)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = Pt(15)

    for paragraph_text in cover_letter_text.split("\n\n"):
        text = paragraph_text.strip()
        if text:
            doc.add_paragraph(text)

    output_filename = f"cover-letter-{uuid.uuid4()}.docx"
    output_path = settings.output_dir / output_filename
    doc.save(str(output_path))
    return output_filename


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


async def generate_draft(
    *,
    core_skills: list[str],
    resume_text: str = "",
    job_title: str,
    organization: str,
    job_description: str,
    cover_letter_sentiment: str | None,
    history_context: str = "",
) -> dict[str, Any]:
    """
    Phase 1 of human-in-the-loop: run LLM chains and return the raw draft
    content for user review.  No docx is generated and nothing is persisted.

    Returns a dict with keys: summary, experiences (list), skills, education,
    certifications, cover_letter, original_resume_text.
    """
    # ---- Resume chain ----
    resume_parser = PydanticOutputParser(pydantic_object=TailoredContent)
    resume_prompt = ChatPromptTemplate.from_messages(
        [("system", _RESUME_SYSTEM), ("human", _RESUME_HUMAN)]
    )

    resume_llm = ChatOpenAI(
        model=settings.openai_model,
        temperature=0.4,
        max_tokens=4096,
        api_key=settings.openai_api_key,
    )

    resume_chain = resume_prompt | resume_llm | resume_parser

    tailored: TailoredContent = await resume_chain.ainvoke(
        {
            "core_skills": ", ".join(core_skills),
            "resume_text": resume_text,
            "history_context": history_context,
            "format_instructions": resume_parser.get_format_instructions(),
            "job_title": job_title,
            "organization": organization,
            "job_description": job_description,
        }
    )

    # ---- Cover-letter chain ----
    cl_parser = PydanticOutputParser(pydantic_object=CoverLetterContent)
    sentiment_key = (cover_letter_sentiment or "formal").lower().strip()
    sentiment_preamble = _SENTIMENT_PREAMBLES.get(
        sentiment_key, _SENTIMENT_PREAMBLES["formal"]
    )

    cl_prompt = ChatPromptTemplate.from_messages(
        [("system", _COVER_LETTER_SYSTEM), ("human", _COVER_LETTER_HUMAN)]
    )

    cl_llm = ChatOpenAI(
        model=settings.openai_model,
        temperature=0.7,
        max_tokens=2048,
        api_key=settings.openai_api_key,
    )

    cl_chain = cl_prompt | cl_llm | cl_parser

    cover_letter: CoverLetterContent = await cl_chain.ainvoke(
        {
            "sentiment_preamble": sentiment_preamble,
            "core_skills": ", ".join(core_skills),
            "resume_text": resume_text,
            "job_title": job_title,
            "organization": organization,
            "format_instructions": cl_parser.get_format_instructions(),
            "job_description": job_description,
        }
    )

    return {
        **tailored.model_dump(),
        "cover_letter": cover_letter.cover_letter,
        "original_resume_text": resume_text,
    }


# ---------------------------------------------------------------------------
# Single-section regeneration
# ---------------------------------------------------------------------------

_SECTION_REGEN_SYSTEM = (
    "You are a senior resume-tailoring specialist. The user has already generated "
    "a full tailored resume. They now want to REGENERATE a SINGLE section to "
    "improve it.\n\n"
    "Candidate core skills: {core_skills}.\n"
    "Target role: {job_title} at {organization}.\n\n"
    "The candidate's ACTUAL resume:\n"
    "--- BEGIN CANDIDATE RESUME ---\n"
    "{resume_text}\n"
    "--- END CANDIDATE RESUME ---\n\n"
    "Job description:\n{job_description}\n\n"
    "═══════════════════════════════════════════════════\n"
    "TASK: Regenerate ONLY the '{section_type}' section.\n"
    "═══════════════════════════════════════════════════\n"
    "{section_instructions}\n\n"
    "The PREVIOUS version of this section was:\n"
    "--- BEGIN CURRENT CONTENT ---\n"
    "{current_content}\n"
    "--- END CURRENT CONTENT ---\n\n"
    "Generate an IMPROVED version that is DIFFERENT from the previous one. "
    "Make it more detailed, better aligned with the JD, or better worded. "
    "Return ONLY the new content for this section as plain text (no JSON wrapper)."
)

_SECTION_INSTRUCTIONS: dict[str, str] = {
    "summary": (
        "Write a 4-5 sentence professional summary that bridges the "
        "candidate's REAL background with the target JD. Highlight strengths "
        "matching the JD priorities. Do NOT claim expertise not in the resume."
    ),
    "experience": (
        "Regenerate this single experience entry. Keep the REAL job title, "
        "organization, and dates from the resume. Produce 4-7 detailed "
        "achievement bullets emphasizing alignment with the JD. Use the "
        "format: 'Title | Organization | Dates\\n• bullet 1\\n• bullet 2...' "
        "Never fabricate achievements not in the resume."
    ),
    "skills": (
        "Produce a comprehensive comma-separated list of skills EXTRACTED "
        "from the resume, reordered by relevance to the JD. Include "
        "technical skills, tools, methodologies, and soft skills."
    ),
    "education": (
        "Extract ALL education entries from the resume. Include degree(s), "
        "institution(s), date(s), honours, and relevant coursework."
    ),
    "certifications": (
        "Extract ALL certifications, licenses, and professional development "
        "entries from the resume."
    ),
    "cover_letter": (
        "Write a compelling, detailed 5-7 paragraph cover letter. "
        "Include a hook, 2-3 body paragraphs with specific achievements "
        "tied to the JD using the CAR method, a values paragraph, and "
        "a confident close. Must read as humanised, not robotic."
    ),
}


async def regenerate_section(
    *,
    core_skills: list[str],
    resume_text: str,
    job_title: str,
    organization: str,
    job_description: str,
    section_id: str,
    current_content: str,
    cover_letter_sentiment: str | None = None,
) -> str:
    """Regenerate a single section of the resume or cover letter."""
    if section_id.startswith("experience_"):
        section_type = "experience"
    else:
        section_type = section_id

    instructions = _SECTION_INSTRUCTIONS.get(
        section_type, _SECTION_INSTRUCTIONS["summary"]
    )

    if section_type == "cover_letter":
        sentiment_key = (cover_letter_sentiment or "formal").lower().strip()
        preamble = _SENTIMENT_PREAMBLES.get(
            sentiment_key, _SENTIMENT_PREAMBLES["formal"]
        )
        instructions = f"{preamble}\n\n{instructions}"

    prompt = ChatPromptTemplate.from_messages([
        ("system", _SECTION_REGEN_SYSTEM),
        ("human", "Regenerate the section now."),
    ])

    llm = ChatOpenAI(
        model=settings.openai_model,
        temperature=0.7,
        max_tokens=2000,
        api_key=settings.openai_api_key,
    )

    chain = prompt | llm
    response = await chain.ainvoke({
        "core_skills": ", ".join(core_skills),
        "resume_text": resume_text,
        "job_title": job_title,
        "organization": organization,
        "job_description": job_description,
        "section_type": section_type,
        "section_instructions": instructions,
        "current_content": current_content,
    })

    return response.content.strip()


def finalize_document(
    *,
    template_path: Path | None,
    content: dict[str, Any],
) -> dict[str, str]:
    """
    Phase 2 of human-in-the-loop: take the (possibly user-edited) content
    and produce the final docx files.

    If *template_path* is provided, inject into the user's template.
    Otherwise build a clean resume docx from scratch.

    Always builds a separate cover-letter docx.

    Returns ``{"resume_url": str, "cover_letter_url": str}``.
    """
    if template_path and template_path.exists():
        resume_fields: dict[str, str] = {}
        for key in ("summary", "skills", "education", "certifications"):
            if key in content and content[key]:
                resume_fields[key] = content[key]
        experiences = content.get("experiences", [])
        if isinstance(experiences, list):
            for i, exp in enumerate(experiences, 1):
                resume_fields[f"experience_{i}"] = exp
        else:
            resume_fields["experience_1"] = str(experiences)
        resume_filename = inject_into_template(template_path, resume_fields)
    else:
        resume_filename = build_resume_docx(content)

    cover_letter_text = content.get("cover_letter", "")
    cl_filename = build_cover_letter_docx(cover_letter_text) if cover_letter_text else ""

    return {"resume_url": resume_filename, "cover_letter_url": cl_filename}


async def tailor_resume(
    *,
    core_skills: list[str],
    resume_text: str = "",
    job_title: str,
    organization: str,
    job_description: str,
    cover_letter_sentiment: str | None,
    template_path: Path | None,
    history_context: str = "",
) -> dict[str, str]:
    """
    Full pipeline (kept for backward-compat with clone flow):
      generate_draft → finalize_document in one shot.

    Returns ``{"tailored_resume_url": str, "cover_letter_text": str,
    "cover_letter_url": str}``.
    """
    draft = await generate_draft(
        core_skills=core_skills,
        resume_text=resume_text,
        job_title=job_title,
        organization=organization,
        job_description=job_description,
        cover_letter_sentiment=cover_letter_sentiment,
        history_context=history_context,
    )

    result = finalize_document(
        template_path=template_path,
        content=draft,
    )

    return {
        "tailored_resume_url": result["resume_url"],
        "cover_letter_text": draft["cover_letter"],
        "cover_letter_url": result["cover_letter_url"],
    }
